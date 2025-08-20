import os

import jdatetime
from PyQt5 import uic
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QCheckBox, QComboBox, QFileDialog, QLineEdit, QSpinBox, QWidget
from docx import Document

from controllers.tender_application.project_session_controller import ProjectSession
from controllers.user_session_controller import UserSession
from views.message_box_view import show_message


class ElectricalTab(QWidget):
    def __init__(self, main_view):
        super().__init__()
        uic.loadUi("ui/tender_application/electrical_tab.ui", self)
        self.main_view = main_view

        self.current_project = ProjectSession()
        self.electrical_specs = self.current_project.project_electrical_specs

        self.set_electrical_ui_values()

        self._initialize_components()

    def _initialize_components(self):
        """Initialize all UI components with their event handlers"""

        """ ------------ Bagfilter ------------ """
        self.bagfilter_type.currentIndexChanged.connect(self._handle_bagfilter_type_changed)
        self.bagfilter_order.textChanged.connect(self._handle_bagfilter_order_changed)
        self.plc_series.currentIndexChanged.connect(self._handle_plc_series_changed)
        self.plc_protocol.currentIndexChanged.connect(self._handle_plc_protocol_changed)
        self.touch_panel_model.currentIndexChanged.connect(self._handle_touch_panel_changed)
        self.olm.stateChanged.connect(self._handle_olm_changed)
        self.ee.stateChanged.connect(self._handle_ee_changed)
        self.me.stateChanged.connect(self._handle_me_changed)
        self.cable_supply.stateChanged.connect(self._handle_cable_supply_changed)
        self.spare_pins.valueChanged.connect(self._handle_spare_pins_changed)

        """ ------------Bagfilter instruments ------------ """
        self.bagfilter_dpt_qty.valueChanged.connect(self._handle_bagfilter_dpt_qty_changed)
        self.bagfilter_dpt_brand.currentIndexChanged.connect(self._handle_bagfilter_dpt_brand_changed)

        self.bagfilter_dps_qty.valueChanged.connect(self._handle_bagfilter_dps_qty_changed)
        self.bagfilter_dps_brand.currentIndexChanged.connect(self._handle_bagfilter_dps_brand_changed)

        self.bagfilter_pt_qty.valueChanged.connect(self._handle_bagfilter_pt_qty_changed)
        self.bagfilter_pt_brand.currentIndexChanged.connect(self._handle_bagfilter_pt_brand_changed)

        self.bagfilter_ps_qty.valueChanged.connect(self._handle_bagfilter_ps_qty_changed)
        self.bagfilter_ps_brand.currentIndexChanged.connect(self._handle_bagfilter_ps_brand_changed)

        self.bagfilter_pg_qty.valueChanged.connect(self._handle_bagfilter_pg_qty_changed)
        self.bagfilter_pg_brand.currentIndexChanged.connect(self._handle_bagfilter_pg_brand_changed)

        self.bagfilter_inlet_tt_qty.valueChanged.connect(self._handle_bagfilter_inlet_tt_qty_changed)
        self.bagfilter_inlet_tt_brand.currentIndexChanged.connect(self._handle_bagfilter_inlet_tt_brand_changed)

        self.bagfilter_outlet_tt_qty.valueChanged.connect(self._handle_bagfilter_outlet_tt_qty_changed)
        self.bagfilter_outlet_tt_brand.currentIndexChanged.connect(self._handle_bagfilter_outlet_tt_brand_changed)

        """ ------------ Transport ------------ """
        self.transport_checkbox.stateChanged.connect(self._handle_transport_checkbox_changed)
        self.screw2_add_btn.clicked.connect(self._show_screw2)
        self.screw2_minuse_btn.clicked.connect(self._hide_screw2)
        self._hide_screw2()

        """ ------------Transport motors ------------ """
        self.rotary_qty.valueChanged.connect(self._handle_rotary_qty_changed)
        self.rotary_kw.currentIndexChanged.connect(self._handle_rotary_kw_changed)

        self.telescopic_chute_qty.valueChanged.connect(self._handle_telescopic_chute_qty_changed)
        self.telescopic_chute_kw.currentIndexChanged.connect(self._handle_telescopic_chute_kw_changed)

        self.slide_gate_qty.valueChanged.connect(self._handle_slide_gate_qty_changed)
        self.slide_gate_kw.currentIndexChanged.connect(self._handle_slide_gate_kw_changed)

        self.screw1_qty.valueChanged.connect(self._handle_screw1_qty_changed)
        self.screw1_kw.currentIndexChanged.connect(self._handle_screw1_kw_changed)

        self.screw2_qty.valueChanged.connect(self._handle_screw2_qty_changed)
        self.screw2_kw.currentIndexChanged.connect(self._handle_screw2_kw_changed)

        """ ------------Transport instruments ------------ """
        self.transport_spd_qty.valueChanged.connect(self._handle_transport_spd_qty_changed)
        self.transport_spd_brand.currentIndexChanged.connect(self._handle_transport_spd_brand_changed)

        self.transport_zs_qty.valueChanged.connect(self._handle_transport_zs_qty_changed)
        self.transport_zs_brand.currentIndexChanged.connect(self._handle_transport_zs_brand_changed)

        self.transport_ls_bin_qty.valueChanged.connect(self._handle_transport_ls_bin_qty_changed)
        self.transport_ls_bin_brand.currentIndexChanged.connect(self._handle_transport_ls_bin_brand_changed)

        self.transport_lt_qty.valueChanged.connect(self._handle_transport_lt_qty_changed)
        self.transport_lt_brand.currentIndexChanged.connect(self._handle_transport_lt_brand_changed)

        """ ------------ Vibration ------------ """
        self.vibration_checkbox.stateChanged.connect(self._handle_vibration_checkbox_changed)

        """ ------------Vibration motors ------------ """
        self.vibration_motor_qty.valueChanged.connect(self._handle_vibration_motor_qty_changed)
        self.vibration_motor_kw.currentIndexChanged.connect(self._handle_vibration_motor_kw_changed)

        """ ------------ Fresh Air ------------ """
        self.freshair_checkbox.stateChanged.connect(self._handle_freshair_checkbox_changed)

        """ ------------Fresh air motors ------------ """
        self.freshair_motor_qty.valueChanged.connect(self._handle_freshair_motor_qty_changed)
        self.freshair_motor_kw.currentIndexChanged.connect(self._handle_freshair_motor_kw_changed)
        self.freshair_motor_start_type.currentIndexChanged.connect(self._handle_freshair_motor_start_type_changed)

        self.freshair_flap_motor_qty.valueChanged.connect(self._handle_freshair_flap_motor_qty_changed)
        self.freshair_flap_motor_kw.currentIndexChanged.connect(self._handle_freshair_flap_motor_kw_changed)
        self.freshair_flap_motor_start_type.currentIndexChanged.connect(self._handle_freshair_flap_start_type_changed)

        self.emergency_flap_motor_qty.valueChanged.connect(self._handle_emergency_flap_motor_qty_changed)
        self.emergency_flap_motor_kw.currentIndexChanged.connect(self._handle_emergency_flap_motor_kw_changed)
        self.emergency_flap_motor_start_type.currentIndexChanged.connect(self._handle_emergency_flap_start_type_changed)

        """ ------------Fresh air instruments ------------ """
        self.freshair_tt_qty.valueChanged.connect(self._handle_freshair_tt_qty_changed)
        self.freshair_tt_brand.currentIndexChanged.connect(self._handle_freshair_tt_brand_changed)

        self.freshair_zs_qty.valueChanged.connect(self._handle_freshair_zs_qty_changed)
        self.freshair_zs_brand.currentIndexChanged.connect(self._handle_freshair_zs_brand_changed)

        """ ------------ Heater Hopper ------------ """
        self.hopper_heater_checkbox.stateChanged.connect(self._handle_hopper_heater_checkbox_changed)

        self.hopper_heater_qty.valueChanged.connect(self._handle_hopper_heater_qty_changed)
        self.hopper_heater_kw.currentIndexChanged.connect(self._handle_hopper_heater_kw_changed)
        self.hopper_heater_ptc_brand.currentIndexChanged.connect(self._handle_hopper_heater_ptc_brand_changed)

        """ ------------ Damper ------------ """
        self.damper_checkbox.stateChanged.connect(self._handle_damper_checkbox_changed)

        self.damper_qty.valueChanged.connect(self._handle_damper_qty_changed)
        self.damper_kw.currentIndexChanged.connect(self._handle_damper_kw_changed)
        self.damper_start_type.currentIndexChanged.connect(self._handle_damper_start_type_changed)

        self.damper_zs_qty.valueChanged.connect(self._handle_damper_zs_qty_changed)
        self.damper_zs_brand.currentIndexChanged.connect(self._handle_damper_zs_brand_changed)

        """ ------------ Fan ------------ """
        self.fan_checkbox.stateChanged.connect(self._handle_fan_checkbox_changed)
        self.qss_btn.clicked.connect(self._handle_qss_generate)

        """ ------------Fan motor properties ------------ """
        self.fan_kw.currentIndexChanged.connect(self._handle_fan_kw_changed)
        self.fan_rpm.textChanged.connect(self._handle_fan_rpm_changed)
        self.fan_start_type.currentIndexChanged.connect(self._handle_fan_start_type_changed)
        self.fan_brand.currentIndexChanged.connect(self._handle_fan_brand_changed)
        self.fan_cooling_method.currentIndexChanged.connect(self._handle_fan_cooling_method_changed)
        self.fan_ip.currentIndexChanged.connect(self._handle_fan_ip_changed)
        self.fan_efficiency_class.currentIndexChanged.connect(self._handle_fan_efficiency_class_changed)
        self.fan_voltage_type.currentIndexChanged.connect(self._handle_fan_voltage_type_changed)
        self.fan_painting_ral.currentIndexChanged.connect(self._handle_fan_painting_ral_changed)
        self.fan_thermal_protection.currentIndexChanged.connect(self._handle_fan_thermal_protection_changed)
        self.fan_space_heater.stateChanged.connect(self._handle_fan_space_heater_changed)
        self.fan_de_nde.stateChanged.connect(self._handle_fan_de_nde_changed)

        """ ------------Fan instruments ------------ """
        self.pt100_qty.valueChanged.connect(self._handle_fan_pt100_qty_changed)
        self.pt100_brand.currentIndexChanged.connect(self._handle_fan_pt100_brand_changed)

        self.fan_bearing_vt_qty.valueChanged.connect(self._handle_fan_bearing_vt_qty_changed)
        self.fan_bearing_vt_brand.currentIndexChanged.connect(self._handle_fan_bearing_vt_brand_changed)

        self.fan_pt_qty.valueChanged.connect(self._handle_fan_pt_qty_changed)
        self.fan_pt_brand.currentIndexChanged.connect(self._handle_fan_pt_brand_changed)

        self.fan_tt_qty.valueChanged.connect(self._handle_fan_tt_qty_changed)
        self.fan_tt_brand.currentIndexChanged.connect(self._handle_fan_tt_brand_changed)

    """ ------------Transport motor handlers ------------ """

    def _handle_rotary_qty_changed(self, value):
        self._handle_rotary_screw_qty(self.rotary_qty)

    def _handle_rotary_kw_changed(self):
        self._handle_combobox_float_kilo(["transport", "motors", "rotary", "power"])

    def _handle_telescopic_chute_qty_changed(self, value):
        self._update_project_value(["transport", "motors", "telescopic_chute", "qty"], value)

    def _handle_telescopic_chute_kw_changed(self):
        self._handle_combobox_float_kilo(["transport", "motors", "telescopic_chute", "power"])

    def _handle_slide_gate_qty_changed(self, value):
        self._update_project_value(["transport", "motors", "slide_gate", "qty"], value)
        """ ------------Update proximity switch quantity based on slide gate qty ------------ """
        if value != 0:
            required_switches = value * 2
            if self.transport_zs_qty.value() < required_switches:
                self.transport_zs_qty.setValue(required_switches)
                self.electrical_specs["transport"]["instruments"]["proximity_switch"]["qty"] = required_switches
        else:
            self.transport_zs_qty.setValue(0)
            self.electrical_specs["transport"]["instruments"]["proximity_switch"]["qty"] = 0

    def _handle_slide_gate_kw_changed(self):
        self._handle_combobox_float_kilo(["transport", "motors", "slide_gate", "power"])

    def _handle_screw1_qty_changed(self, value):
        self._handle_rotary_screw_qty(self.screw1_qty)

    def _handle_screw1_kw_changed(self):
        self._handle_combobox_float_kilo(["transport", "motors", "screw1", "power"])

    def _handle_screw2_qty_changed(self, value):
        self._handle_rotary_screw_qty(self.screw2_qty)

    def _handle_screw2_kw_changed(self):
        self._handle_combobox_float_kilo(["transport", "motors", "screw2", "power"])

    """ ------------Transport instrument handlers ------------ """

    def _handle_transport_spd_qty_changed(self, value):
        self._update_project_value(["transport", "instruments", "speed_detector", "qty"], value)

    def _handle_transport_spd_brand_changed(self):
        self._update_project_value(["transport", "instruments", "speed_detector", "brand"])

    def _handle_transport_zs_qty_changed(self, value):
        self._update_project_value(["transport", "instruments", "proximity_switch", "qty"], value)

    def _handle_transport_zs_brand_changed(self):
        self._update_project_value(["transport", "instruments", "proximity_switch", "brand"])

    def _handle_transport_ls_bin_qty_changed(self, value):
        self._update_project_value(["transport", "instruments", "level_switch", "qty"], value)

    def _handle_transport_ls_bin_brand_changed(self):
        self._update_project_value(["transport", "instruments", "level_switch", "brand"])

    def _handle_transport_lt_qty_changed(self, value):
        self._update_project_value(["transport", "instruments", "level_transmitter", "qty"], value)

    def _handle_transport_lt_brand_changed(self):
        self._update_project_value(["transport", "instruments", "level_transmitter", "brand"])

    """ ------------Vibration motor handlers ------------ """

    def _handle_vibration_motor_qty_changed(self, value):
        self._update_project_value(["vibration", "motors", "vibration", "qty"], value)

    def _handle_vibration_motor_kw_changed(self):
        self._handle_combobox_float_kilo(["vibration", "motors", "vibration", "power"])

    """ ------------Fresh air motor handlers ------------ """

    def _handle_freshair_motor_qty_changed(self, value):
        self._update_project_value(["fresh_air", "motors", "freshair_motor", "qty"], value)
        """ ------------Update proximity switch quantity based on motor qty ------------ """
        if self.freshair_motor_qty.value() != 0:
            required_switches = self.freshair_motor_qty.value() * 2
            if self.freshair_zs_qty.value() < required_switches:
                self.freshair_zs_qty.setValue(required_switches)
                self.electrical_specs["fresh_air"]["instruments"]["proximity_switch"]["qty"] = required_switches

    def _handle_freshair_motor_kw_changed(self):
        self._handle_combobox_float_kilo(["fresh_air", "motors", "freshair_motor", "power"])

    def _handle_freshair_motor_start_type_changed(self):
        self._update_project_value(["fresh_air", "motors", "freshair_motor", "start_type"])

    def _handle_freshair_flap_motor_qty_changed(self, value):
        self._update_project_value(["fresh_air", "motors", "fresh_air_flap", "qty"], value)

    def _handle_freshair_flap_motor_kw_changed(self):
        self._handle_combobox_float_kilo(["fresh_air", "motors", "fresh_air_flap", "power"])

    def _handle_freshair_flap_start_type_changed(self):
        self._update_project_value(["fresh_air", "motors", "fresh_air_flap", "start_type"])

    def _handle_emergency_flap_motor_qty_changed(self, value):
        self._update_project_value(["fresh_air", "motors", "emergency_flap", "qty"], value)

    def _handle_emergency_flap_motor_kw_changed(self):
        self._handle_combobox_float_kilo(["fresh_air", "motors", "emergency_flap", "power"])

    def _handle_emergency_flap_start_type_changed(self):
        self._update_project_value(["fresh_air", "motors", "emergency_flap", "start_type"])

    """ ------------Fresh air instrument handlers ------------ """

    def _handle_freshair_tt_qty_changed(self, value):
        self._update_project_value(["fresh_air", "instruments", "temperature_transmitter", "qty"], value)

    def _handle_freshair_tt_brand_changed(self):
        self._update_project_value(["fresh_air", "instruments", "temperature_transmitter", "brand"])

    def _handle_freshair_zs_qty_changed(self, value):
        self._update_project_value(["fresh_air", "instruments", "proximity_switch", "qty"], value)

    def _handle_freshair_zs_brand_changed(self):
        self._update_project_value(["fresh_air", "instruments", "proximity_switch", "brand"])

    """ ------------Hopper heater handlers ------------ """

    def _handle_hopper_heater_qty_changed(self, value):
        self._update_project_value(["hopper_heater", "motors", "elements", "qty"], value)

        """ ------------Update PTC qty based on heater qty ------------ """
        ptc_qty = value * 2
        self.electrical_specs["hopper_heater"]["instruments"]["ptc"]["qty"] = ptc_qty
        self.hopper_heater_ptc_qty.setText(f"Qty: {ptc_qty}")

    def _handle_hopper_heater_kw_changed(self):
        self._handle_combobox_float_kilo(["hopper_heater", "motors", "elements", "power"])

    def _handle_hopper_heater_ptc_brand_changed(self):
        self._update_project_value(["hopper_heater", "instruments", "ptc", "brand"])

    """ ------------Damper handlers ------------ """
    def _handle_damper_qty_changed(self, value):
        self._update_project_value(["damper", "motors", "damper", "qty"], value)

    def _handle_damper_kw_changed(self):
        self._handle_combobox_float_kilo(["damper", "motors", "damper", "power"])

    def _handle_damper_start_type_changed(self):
        self._update_project_value(["damper", "motors", "damper", "start_type"])

    def _handle_damper_zs_qty_changed(self, value):
        self._update_project_value(["damper", "instruments", "proximity_switch", "qty"], value)

    def _handle_damper_zs_brand_changed(self):
        self._update_project_value(["damper", "instruments", "proximity_switch", "brand"])

    """ ------------Fan motor handlers ------------ """

    def _handle_qss_generate(self):
        create_qss_word()

    def _handle_fan_kw_changed(self):
        self._handle_combobox_float_kilo(["fan", "motors", "fan", "power"])

    def _handle_fan_rpm_changed(self, text):
        self._update_project_value(["fan", "motors", "fan", "rpm"])

    def _handle_fan_start_type_changed(self):
        self._update_project_value(["fan", "motors", "fan", "start_type"])
        if self.fan_start_type.currentText() == "VFD" or self.fan_start_type.currentText() == "Soft Starter":
            self.fan_voltage_type.setCurrentText("LV")
            self.fan_voltage_type.setEnabled(False)
            self._update_project_value(["fan", "motors", "fan", "voltage_type"], "LV")
            self.fan_voltage_type_text.setText("Voltage Type(LV)")
        else:
            self.fan_voltage_type.setCurrentIndex(0)
            self.fan_voltage_type.setEnabled(True)
            self._update_project_value(["fan", "motors", "fan", "voltage_type"], self.fan_voltage_type.itemText(0))
            self.fan_voltage_type_text.setText("Voltage Type")

    def _handle_fan_brand_changed(self):
        self._update_project_value(["fan", "motors", "fan", "brand"])

    def _handle_fan_cooling_method_changed(self):
        self._update_project_value(["fan", "motors", "fan", "cooling_method"])

    def _handle_fan_ip_changed(self):
        self._update_project_value(["fan", "motors", "fan", "ip_rating"])

    def _handle_fan_efficiency_class_changed(self):
        self._update_project_value(["fan", "motors", "fan", "efficiency_class"])

    def _handle_fan_voltage_type_changed(self):
        self._update_project_value(["fan", "motors", "fan", "voltage_type"])

    def _handle_fan_painting_ral_changed(self):
        self._update_project_value(["fan", "motors", "fan", "painting_ral"])

    def _handle_fan_thermal_protection_changed(self):
        self._update_project_value(["fan", "motors", "fan", "thermal_protection"])

    def _handle_fan_space_heater_changed(self, state):
        self._handle_checkbox(["fan", "motors", "fan", "space_heater"])

    def _handle_fan_de_nde_changed(self, state):
        self._handle_checkbox(["fan", "motors", "fan", "de_nde"])
        pt100_value = self.pt100_qty.value()

        if state:
            self.pt100_qty.setValue(pt100_value + 2)
            self._update_project_value(["fan", "instruments", "pt100", "qty"], pt100_value + 2)

        else:
            self.pt100_qty.setValue(pt100_value - 2)
            self._update_project_value(["fan", "instruments", "pt100", "qty"], max(pt100_value - 2, 0))

    """ ------------Fan instrument handlers ------------ """

    def _handle_fan_pt100_qty_changed(self, value):
        self._update_project_value(["fan", "instruments", "pt100", "qty"], value)

    def _handle_fan_pt100_brand_changed(self):
        self._update_project_value(["fan", "instruments", "pt100", "brand"])

    def _handle_fan_bearing_vt_qty_changed(self, value):
        self._update_project_value(["fan", "instruments", "bearing_vibration_transmitter", "qty"], value)

    def _handle_fan_bearing_vt_brand_changed(self):
        self._update_project_value(["fan", "instruments", "bearing_vibration_transmitter", "brand"])

    def _handle_fan_pt_qty_changed(self, value):
        self._update_project_value(["fan", "instruments", "pressure_transmitter", "qty"], value)

    def _handle_fan_pt_brand_changed(self):
        self._update_project_value(["fan", "instruments", "pressure_transmitter", "brand"])

    def _handle_fan_tt_qty_changed(self, value):
        self._update_project_value(["fan", "instruments", "temperature_transmitter", "qty"], value)

    def _handle_fan_tt_brand_changed(self):
        self._update_project_value(["fan", "instruments", "temperature_transmitter", "brand"])

    def _handle_rotary_screw_qty(self, widget):
        """Special handler for rotary and screw qty fields that affect speed detector"""
        motor_type = None
        if widget == self.rotary_qty:
            motor_type = "rotary"
        elif widget == self.screw1_qty:
            motor_type = "screw1"
        elif widget == self.screw2_qty:
            motor_type = "screw2"

        if not motor_type:
            return

        """ ------------Update the tender_application details ------------ """
        self.electrical_specs["transport"]["motors"][motor_type]["qty"] = widget.value()

        """ ------------Calculate total rotary/screw motors ------------ """
        total_motors = (self.rotary_qty.value() +
                        self.screw1_qty.value() +
                        self.screw2_qty.value())

        """ ------------Update speed detector quantity based on total motors ------------ """
        if total_motors == 0:
            self.transport_spd_qty.setValue(0)
            self.electrical_specs["transport"]["instruments"]["speed_detector"]["qty"] = 0
        elif self.transport_spd_qty.value() < total_motors:
            self.transport_spd_qty.setValue(total_motors)
            self.electrical_specs["transport"]["instruments"]["speed_detector"]["qty"] = total_motors

    def _handle_transport_checkbox_changed(self, state):
        """Handle transport section enabling/disabling"""
        enabled = state == Qt.Checked
        self.electrical_specs["transport"]["status"] = enabled

        if not enabled:
            self._reset_qtys(self.electrical_specs["transport"])

        self._reset_section_widgets(self.transport_gbox, enabled)
        self.transport_checkbox.setEnabled(True)

    def _handle_vibration_checkbox_changed(self, state):
        """Handle vibration section enabling/disabling"""
        enabled = state == Qt.Checked
        self.electrical_specs["vibration"]["status"] = enabled

        if not enabled:
            self._reset_qtys(self.electrical_specs["vibration"])

        self._reset_section_widgets(self.vibration_gbox, enabled)
        self.vibration_checkbox.setEnabled(True)

    def _handle_freshair_checkbox_changed(self, state):
        """Handle fresh air section enabling/disabling"""
        enabled = state == Qt.Checked
        self.electrical_specs["fresh_air"]["status"] = enabled

        if not enabled:
            self._reset_qtys(self.electrical_specs["fresh_air"])

        self._reset_section_widgets(self.freshair_gbox, enabled)
        self.freshair_checkbox.setEnabled(True)

    def _handle_hopper_heater_checkbox_changed(self, state):
        """Handle hopper heater section enabling/disabling"""
        enabled = state == Qt.Checked
        self.electrical_specs["hopper_heater"]["status"] = enabled

        if not enabled:
            self._reset_qtys(self.electrical_specs["hopper_heater"])

        self._reset_section_widgets(self.hopper_heater_gbox, enabled)
        self.hopper_heater_checkbox.setEnabled(True)

    def _handle_damper_checkbox_changed(self, state):
        """Handle damper section enabling/disabling"""
        enabled = state == Qt.Checked
        self.electrical_specs["damper"]["status"] = enabled
        self.electrical_specs["damper"]["motors"]["damper"]["qty"] = 1 if enabled else 0

        if not enabled:
            self._reset_qtys(self.electrical_specs["damper"])

        """ ------------Enable/disable individual damper widgets ------------ """
        self.damper_kw.setEnabled(enabled)
        self.damper_qty.setEnabled(enabled)
        self.damper_start_type.setEnabled(enabled)
        self.damper_zs_qty.setEnabled(enabled)
        self.damper_zs_brand.setEnabled(enabled)

        """ ------------Reset values if disabling ------------ """
        if not enabled:
            self.damper_kw.setCurrentIndex(0)
            self.damper_qty.setValue(0)
            self.damper_start_type.setCurrentIndex(0)
            self.damper_zs_qty.setValue(0)
            self.damper_zs_brand.setCurrentIndex(0)

    def _handle_fan_checkbox_changed(self, state):
        """Handle fan section enabling/disabling"""
        enabled = state == Qt.Checked
        self.electrical_specs["fan"]["status"] = enabled
        self.electrical_specs["fan"]["motors"]["fan"]["qty"] = 1 if enabled else 0

        if not enabled:
            self._reset_qtys(self.electrical_specs["fan"])

        """ ------------Enable/disable individual fan widgets ------------ """
        fan_widgets = [
            self.fan_kw,
            self.fan_rpm,
            self.fan_start_type,
            self.fan_brand,
            self.fan_cooling_method,
            self.fan_ip,
            self.fan_efficiency_class,
            self.fan_voltage_type,
            self.fan_painting_ral,
            self.fan_thermal_protection,
            self.fan_space_heater,
            self.fan_bearing_vt_qty,
            self.fan_bearing_vt_brand,
            self.fan_pt_qty,
            self.fan_pt_brand,
            self.fan_tt_qty,
            self.fan_tt_brand,
            self.fan_de_nde,
            self.pt100_qty,
            self.pt100_brand,
            self.qss_btn
        ]

        """ ------------Enable/disable all fan widgets ------------ """
        for widget in fan_widgets:
            widget.setEnabled(enabled)

        """ ------------Reset values if disabling ------------ """
        if not enabled:
            self.fan_kw.setCurrentIndex(0)
            self.fan_rpm.setText("")
            self.fan_start_type.setCurrentIndex(0)
            self.fan_brand.setCurrentIndex(0)
            self.fan_cooling_method.setCurrentIndex(0)
            self.fan_ip.setCurrentIndex(0)
            self.fan_efficiency_class.setCurrentIndex(0)
            self.fan_voltage_type.setCurrentIndex(0)
            self.fan_painting_ral.setCurrentIndex(0)
            self.fan_thermal_protection.setCurrentIndex(0)
            self.fan_de_nde.setChecked(False)
            self.fan_space_heater.setChecked(False)
            """ ------------Reset instrument values ------------ """
            for widget in fan_widgets[11:]:  # All widgets after fan_space_heater
                if isinstance(widget, QSpinBox):
                    widget.setValue(0)
                elif isinstance(widget, QComboBox):
                    widget.setCurrentIndex(0)

        self.fan_checkbox.setEnabled(True)

    def _show_screw2(self):
        """Show screw2 related widgets"""
        self.screw2_label.setVisible(True)
        self.screw2_minuse_btn.setVisible(True)
        self.screw2_qty.setVisible(True)
        self.screw2_kw.setVisible(True)

    def _hide_screw2(self):
        """Hide screw2 related widgets and reset values"""
        self.screw2_label.setVisible(False)
        self.screw2_minuse_btn.setVisible(False)
        self.screw2_qty.setVisible(False)
        self.screw2_kw.setVisible(False)
        self.screw2_qty.setValue(0)
        self.screw2_kw.setCurrentIndex(0)

    def _handle_bagfilter_type_changed(self):
        if self.bagfilter_type.currentIndex() == 0:
            self.bagfilter_order.setPlaceholderText("Order")
        if self.bagfilter_type.currentText() == "Griin/China":  # Griin/China
            self.bagfilter_order.setPlaceholderText("Ex: 8.96×5.(2.7m).10")  # 5 valve ~ compartment
        if self.bagfilter_type.currentText() == "BETH":  # BETH
            self.bagfilter_order.setPlaceholderText("Ex: 6.78x2.3.10")  # 6x2 valve

        self._update_project_value(["bagfilter", "type"], self.bagfilter_type.currentText())

    def _handle_bagfilter_order_changed(self, value):
        self._update_project_value(["bagfilter", "order"], value)

    def _handle_plc_series_changed(self, index):
        self._update_project_value(["bagfilter", "plc_series"], self.plc_series.currentText())

    def _handle_plc_protocol_changed(self):
        self._update_project_value(["bagfilter", "plc_protocol"])
        if self.plc_protocol.currentText() == "PROFIBUS":
            self.olm.setEnabled(True)
        else:
            self.olm.setEnabled(False)
            self.olm.setChecked(False)
            self._update_project_value(["bagfilter", "olm"], False)

    def _handle_touch_panel_changed(self):
        if self.touch_panel_model.currentText() == "None":
            self._update_project_value(["bagfilter", "touch_panel"], "")
        else:
            self._update_project_value(["bagfilter", "touch_panel"], self.touch_panel_model.currentText())

    def _handle_olm_changed(self, state):
        self._update_project_value(["bagfilter", "olm"], state == Qt.Checked)

    def _handle_ee_changed(self, state):
        self._update_project_value(["bagfilter", "ee"], state == Qt.Checked)

    def _handle_me_changed(self, state):
        self._update_project_value(["bagfilter", "me"], state == Qt.Checked)

    def _handle_cable_supply_changed(self, state):
        self._update_project_value(["bagfilter", "cable_supply"], state == Qt.Checked)

    def _handle_cable_length_changed(self, value):
        self._update_project_value(["bagfilter", "cable_dimension"], value)

    """ ------------ Bagfilter instrument handlers ------------ """

    def _handle_spare_pins_changed(self, value):
        self._update_project_value(["bagfilter", "spare_pins"], value)

    def _handle_bagfilter_dpt_qty_changed(self, value):
        self._update_project_value(["bagfilter", "instruments", "delta_pressure_transmitter", "qty"], value)

    def _handle_bagfilter_dpt_brand_changed(self, index):
        self._update_project_value(["bagfilter", "instruments", "delta_pressure_transmitter", "brand"],
                                   self.bagfilter_dpt_brand.currentText())

    def _handle_bagfilter_dps_qty_changed(self, value):
        self._update_project_value(["bagfilter", "instruments", "delta_pressure_switch", "qty"], value)

    def _handle_bagfilter_dps_brand_changed(self, index):
        self._update_project_value(["bagfilter", "instruments", "delta_pressure_switch", "brand"],
                                   self.bagfilter_dps_brand.currentText())

    def _handle_bagfilter_pt_qty_changed(self, value):
        self._update_project_value(["bagfilter", "instruments", "pressure_transmitter", "qty"], value)

    def _handle_bagfilter_pt_brand_changed(self, index):
        self._update_project_value(["bagfilter", "instruments", "pressure_transmitter", "brand"],
                                   self.bagfilter_pt_brand.currentText())

    def _handle_bagfilter_ps_qty_changed(self, value):
        self._update_project_value(["bagfilter", "instruments", "pressure_switch", "qty"], value)

    def _handle_bagfilter_ps_brand_changed(self, index):
        self._update_project_value(["bagfilter", "instruments", "pressure_switch", "brand"],
                                   self.bagfilter_ps_brand.currentText())

    def _handle_bagfilter_pg_qty_changed(self, value):
        self._update_project_value(["bagfilter", "instruments", "pressure_gauge", "qty"], value)

    def _handle_bagfilter_pg_brand_changed(self, index):
        self._update_project_value(["bagfilter", "instruments", "pressure_gauge", "brand"],
                                   self.bagfilter_pg_brand.currentText())

    def _handle_bagfilter_inlet_tt_qty_changed(self, value):
        self._update_project_value(["bagfilter", "instruments", "inlet_temperature_transmitter", "qty"], value)

    def _handle_bagfilter_inlet_tt_brand_changed(self, index):
        self._update_project_value(["bagfilter", "instruments", "inlet_temperature_transmitter", "brand"],
                                   self.bagfilter_inlet_tt_brand.currentText())

    def _handle_bagfilter_outlet_tt_qty_changed(self, value):
        self._update_project_value(["bagfilter", "instruments", "outlet_temperature_transmitter", "qty"], value)

    def _handle_bagfilter_outlet_tt_brand_changed(self, index):
        self._update_project_value(["bagfilter", "instruments", "outlet_temperature_transmitter", "brand"],
                                   self.bagfilter_outlet_tt_brand.currentText())

    def _update_project_value(self, path_list, value=None):
        """Update tender_application details dictionary value at the specified path

        Args:
            path_list: List of keys to navigate the nested dictionary
            value: Value to set (if None, gets from sender widget)
        """
        if value is None:
            """ ------------Get value from sender ------------ """
            sender = self.sender()
            if isinstance(sender, QComboBox):
                value = sender.currentText()
            elif isinstance(sender, QSpinBox):
                value = sender.value()
            elif isinstance(sender, QLineEdit):
                value = sender.text()
            elif isinstance(sender, QCheckBox):
                value = sender.isChecked()
            else:
                return

        """ ------------Navigate to the right location in the dictionary ------------ """
        target = self.electrical_specs
        for key in path_list[:-1]:
            target = target[key]

        """ ------------Set the value ------------ """
        target[path_list[-1]] = value

    def _handle_combobox_float_kilo(self, path_list):

        sender = self.sender()
        text = sender.currentText()
        try:
            value = float(text) * 1000
        except ValueError:
            value = 0.0

        self._update_project_value(path_list, value)

    def _handle_checkbox(self, path_list):
        """Handle checkbox state change

        Args:
            path_list: List of keys to navigate the nested dictionary
        """
        sender = self.sender()
        state = sender.isChecked()
        self._update_project_value(path_list, state)

    def _reset_qtys(self, d):
        """Reset all quantity values in a dictionary to zero

        Args:
            d: Dictionary to process
        """
        for key, value in d.items():
            if isinstance(value, dict):
                self._reset_qtys(value)
            elif key == "qty":
                d[key] = 0

    def _reset_section_widgets(self, parent_widget, enabled=True):
        """Reset and enable/disable all widgets in a section

        Args:
            parent_widget: Parent widget containing child widgets to reset
            enabled: Whether to enable or disable the widgets
        """
        for child in parent_widget.findChildren(QWidget):
            """ ------------Skip the checkbox itself ------------ """
            if isinstance(child, QCheckBox) and child.parent() == parent_widget:
                continue

            """ ------------Reset values if disabling ------------ """
            if not enabled:
                if isinstance(child, QSpinBox):
                    child.setValue(0)
                elif isinstance(child, QComboBox):
                    child.setCurrentIndex(0)
                elif isinstance(child, QLineEdit):
                    child.setText("")
                elif isinstance(child, QCheckBox):
                    child.setChecked(False)

            """ ------------Set enabled state ------------ """
            child.setEnabled(enabled)

    """ ------------ Check Ui Rules ------------ """

    def check_electrical_tab_ui_rules(self):

        if self.bagfilter_order.text == "":
            show_message("Enter Bagfilter Order", "Error")
            return False

        # if self.bagfilter_type.currentIndex() == 1:
        #     griin_pattern = r"^\d+(\.\d+)?x\d+\.?\(\d+(\.\d+)?m\)\.\d+$"
        #     match = re.fullmatch(griin_pattern, self.electrical_specs["bagfilter"]["order"])
        #     if not match:  # Griin/China
        #         show_message("Please Follow Pattern Like 8.96×5.(2.7m).10 for Griin/China Model", "Error")
        #         return False
        # if self.bagfilter_type.currentIndex() == 2:
        #     beth_pattern = r"^(\d+)\.\d+x(\d+)\.(\d+)\.(\d+)$"
        #     match = re.fullmatch(beth_pattern, self.electrical_specs["bagfilter"]["order"])
        #     if not match:  # BETH
        #         show_message("Please Follow Pattern Like 6.78x2.3.10 for BETH Model", "Error")
        #         return False

        if self.touch_panel_model.currentIndex() == 0:
            show_message("Select Touch Panel Model")
            return False

        if self.fan_checkbox.isChecked():
            if self.fan_voltage_type.currentIndex() == 0:
                show_message("Select Fan Voltage Type")
                return False

        return True

    """ Load Pervios Revision as need """

    def set_electrical_ui_values(self):  # Using for open pervios project
        """
        Set values for UI elements based on the self.electrical_specs dictionary.
        """
        new_proj = True if self.current_project.revision == None else False
        if new_proj:  # new proj
            self.transport_checkbox.setChecked(False)
            self.damper_checkbox.setChecked(False)
            self.fan_checkbox.setChecked(False)
            self.hopper_heater_checkbox.setChecked(False)
            self.vibration_checkbox.setChecked(False)
            self.freshair_checkbox.setChecked(False)
            return

        try:
            # QLabel elements
            self.hopper_heater_ptc_qty.setText(str(self.electrical_specs['hopper_heater']['instruments']['ptc']['qty']))

            # QSpinBox elements
            self.spare_pins.setValue(self.electrical_specs['bagfilter']['spare_pins'])
            self.hopper_heater_qty.setValue(self.electrical_specs['hopper_heater']['motors']['elements']['qty'])
            self.vibration_motor_qty.setValue(self.electrical_specs['vibration']['motors']['vibration']['qty'])
            self.emergency_flap_motor_qty.setValue(
                self.electrical_specs['fresh_air']['motors']['emergency_flap']['qty'])
            self.freshair_zs_qty.setValue(self.electrical_specs['fresh_air']['instruments']['proximity_switch']['qty'])
            self.freshair_tt_qty.setValue(
                self.electrical_specs['fresh_air']['instruments']['temperature_transmitter']['qty'])
            self.freshair_motor_qty.setValue(self.electrical_specs['fresh_air']['motors']['freshair_motor']['qty'])
            self.freshair_flap_motor_qty.setValue(self.electrical_specs['fresh_air']['motors']['fresh_air_flap']['qty'])
            self.transport_zs_qty.setValue(self.electrical_specs['transport']['instruments']['proximity_switch']['qty'])
            self.rotary_qty.setValue(self.electrical_specs['transport']['motors']['rotary']['qty'])
            self.telescopic_chute_qty.setValue(self.electrical_specs['transport']['motors']['telescopic_chute']['qty'])
            self.transport_ls_bin_qty.setValue(self.electrical_specs['transport']['instruments']['level_switch']['qty'])
            self.transport_lt_qty.setValue(
                self.electrical_specs['transport']['instruments']['level_transmitter']['qty'])
            self.screw2_qty.setValue(self.electrical_specs['transport']['motors']['screw2']['qty'])
            self.screw1_qty.setValue(self.electrical_specs['transport']['motors']['screw1']['qty'])
            self.slide_gate_qty.setValue(self.electrical_specs['transport']['motors']['slide_gate']['qty'])
            self.transport_spd_qty.setValue(self.electrical_specs['transport']['instruments']['speed_detector']['qty'])
            self.fan_bearing_vt_qty.setValue(
                self.electrical_specs['fan']['instruments']['bearing_vibration_transmitter']['qty'])
            self.damper_qty.setValue(self.electrical_specs['damper']['motors']['damper']['qty'])
            self.damper_zs_qty.setValue(self.electrical_specs['damper']['instruments']['proximity_switch']['qty'])
            self.fan_pt_qty.setValue(self.electrical_specs['fan']['instruments']['pressure_transmitter']['qty'])
            self.fan_tt_qty.setValue(self.electrical_specs['fan']['instruments']['temperature_transmitter']['qty'])
            self.pt100_qty.setValue(self.electrical_specs['fan']['instruments']['pt100']['qty'])
            self.bagfilter_dpt_qty.setValue(
                self.electrical_specs['bagfilter']['instruments']['delta_pressure_transmitter']['qty'])
            self.bagfilter_pt_qty.setValue(
                self.electrical_specs['bagfilter']['instruments']['pressure_transmitter']['qty'])
            self.bagfilter_ps_qty.setValue(self.electrical_specs['bagfilter']['instruments']['pressure_switch']['qty'])
            self.bagfilter_inlet_tt_qty.setValue(
                self.electrical_specs['bagfilter']['instruments']['inlet_temperature_transmitter']['qty'])
            self.bagfilter_dps_qty.setValue(
                self.electrical_specs['bagfilter']['instruments']['delta_pressure_switch']['qty'])
            self.bagfilter_outlet_tt_qty.setValue(
                self.electrical_specs['bagfilter']['instruments']['outlet_temperature_transmitter']['qty'])
            self.bagfilter_pg_qty.setValue(self.electrical_specs['bagfilter']['instruments']['pressure_gauge']['qty'])

            # QComboBox elements
            self.hopper_heater_kw.setCurrentText(
                str(self.electrical_specs['hopper_heater']['motors']['elements']['power'] / 1000))
            self.hopper_heater_ptc_brand.setCurrentText(
                str(self.electrical_specs['hopper_heater']['instruments']['ptc']['brand']) if
                self.electrical_specs['hopper_heater']['instruments']['ptc']['brand'] else "")
            self.vibration_motor_kw.setCurrentText(
                str(self.electrical_specs['vibration']['motors']['vibration']['power'] / 1000))
            self.freshair_zs_brand.setCurrentText(
                str(self.electrical_specs['fresh_air']['instruments']['proximity_switch']['brand']) if
                self.electrical_specs['fresh_air']['instruments']['proximity_switch']['brand'] else "")
            self.freshair_tt_brand.setCurrentText(
                str(self.electrical_specs['fresh_air']['instruments']['temperature_transmitter']['brand']) if
                self.electrical_specs['fresh_air']['instruments']['temperature_transmitter']['brand'] else "")
            self.freshair_flap_motor_start_type.setCurrentText(
                str(self.electrical_specs['fresh_air']['motors']['fresh_air_flap']['start_type']) if
                self.electrical_specs['fresh_air']['motors']['fresh_air_flap']['start_type'] else "")
            self.freshair_motor_start_type.setCurrentText(
                str(self.electrical_specs['fresh_air']['motors']['freshair_motor']['start_type']) if
                self.electrical_specs['fresh_air']['motors']['freshair_motor']['start_type'] else "")
            self.freshair_motor_kw.setCurrentText(
                str(self.electrical_specs['fresh_air']['motors']['freshair_motor']['power'] / 1000))
            self.emergency_flap_motor_start_type.setCurrentText(
                str(self.electrical_specs['fresh_air']['motors']['emergency_flap']['start_type']) if
                self.electrical_specs['fresh_air']['motors']['emergency_flap']['start_type'] else "")
            self.freshair_flap_motor_kw.setCurrentText(
                str(self.electrical_specs['fresh_air']['motors']['fresh_air_flap']['power'] / 1000))
            self.emergency_flap_motor_kw.setCurrentText(
                str(self.electrical_specs['fresh_air']['motors']['emergency_flap']['power'] / 1000))
            self.rotary_kw.setCurrentText(str(self.electrical_specs['transport']['motors']['rotary']['power'] / 1000))
            self.transport_zs_brand.setCurrentText(
                str(self.electrical_specs['transport']['instruments']['proximity_switch']['brand']) if
                self.electrical_specs['transport']['instruments']['proximity_switch']['brand'] else "")
            self.telescopic_chute_kw.setCurrentText(
                str(self.electrical_specs['transport']['motors']['telescopic_chute']['power'] / 1000))
            self.transport_ls_bin_brand.setCurrentText(
                str(self.electrical_specs['transport']['instruments']['level_switch']['brand']) if
                self.electrical_specs['transport']['instruments']['level_switch']['brand'] else "")
            self.transport_lt_brand.setCurrentText(
                str(self.electrical_specs['transport']['instruments']['level_transmitter']['brand']) if
                self.electrical_specs['transport']['instruments']['level_transmitter']['brand'] else "")
            self.screw2_kw.setCurrentText(str(self.electrical_specs['transport']['motors']['screw2']['power'] / 1000))
            self.screw1_kw.setCurrentText(str(self.electrical_specs['transport']['motors']['screw1']['power'] / 1000))
            self.slide_gate_kw.setCurrentText(
                str(self.electrical_specs['transport']['motors']['slide_gate']['power'] / 1000))
            self.transport_spd_brand.setCurrentText(
                str(self.electrical_specs['transport']['instruments']['speed_detector']['brand']) if
                self.electrical_specs['transport']['instruments']['speed_detector']['brand'] else "")
            self.fan_brand.setCurrentText(str(self.electrical_specs['fan']['motors']['fan']['brand']) if
                                          self.electrical_specs['fan']['motors']['fan']['brand'] else "")
            self.fan_thermal_protection.setCurrentText(
                str(self.electrical_specs['fan']['motors']['fan']['thermal_protection']) if
                self.electrical_specs['fan']['motors']['fan']['thermal_protection'] else "")
            self.fan_pt_brand.setCurrentText(
                str(self.electrical_specs['fan']['instruments']['pressure_transmitter']['brand']) if
                self.electrical_specs['fan']['instruments']['pressure_transmitter']['brand'] else "")
            self.fan_start_type.setCurrentText(str(self.electrical_specs['fan']['motors']['fan']['start_type']) if
                                               self.electrical_specs['fan']['motors']['fan']['start_type'] else "")
            self.fan_tt_brand.setCurrentText(
                str(self.electrical_specs['fan']['instruments']['temperature_transmitter']['brand']) if
                self.electrical_specs['fan']['instruments']['temperature_transmitter']['brand'] else "")
            self.pt100_brand.setCurrentText(str(self.electrical_specs['fan']['instruments']['pt100']['brand']) if
                                            self.electrical_specs['fan']['instruments']['pt100']['brand'] else "")
            self.damper_zs_brand.setCurrentText(
                str(self.electrical_specs['damper']['instruments']['proximity_switch']['brand']) if
                self.electrical_specs['damper']['instruments']['proximity_switch']['brand'] else "")
            self.fan_bearing_vt_brand.setCurrentText(
                str(self.electrical_specs['fan']['instruments']['bearing_vibration_transmitter']['brand']) if
                self.electrical_specs['fan']['instruments']['bearing_vibration_transmitter']['brand'] else "")
            self.damper_kw.setCurrentText(str(self.electrical_specs['damper']['motors']['damper']['power'] / 1000))
            self.damper_start_type.setCurrentText(
                str(self.electrical_specs['damper']['motors']['damper']['start_type']) if
                self.electrical_specs['damper']['motors']['damper']['start_type'] else "")
            self.fan_painting_ral.setCurrentText(str(self.electrical_specs['fan']['motors']['fan']['painting_ral']) if
                                                 self.electrical_specs['fan']['motors']['fan']['painting_ral'] else "")
            self.fan_kw.setCurrentText(str(self.electrical_specs['fan']['motors']['fan']['power'] / 1000))
            self.fan_efficiency_class.setCurrentText(
                str(self.electrical_specs['fan']['motors']['fan']['efficiency_class']) if
                self.electrical_specs['fan']['motors']['fan']['efficiency_class'] else "")
            self.fan_cooling_method.setCurrentText(
                str(self.electrical_specs['fan']['motors']['fan']['cooling_method']) if
                self.electrical_specs['fan']['motors']['fan']['cooling_method'] else "")
            self.fan_voltage_type.setCurrentText(str(self.electrical_specs['fan']['motors']['fan']['voltage_type']) if
                                                 self.electrical_specs['fan']['motors']['fan']['voltage_type'] else "")
            self.fan_ip.setCurrentText(str(self.electrical_specs['fan']['motors']['fan']['ip_rating']) if
                                       self.electrical_specs['fan']['motors']['fan']['ip_rating'] else "")
            self.bagfilter_pg_brand.setCurrentText(
                str(self.electrical_specs['bagfilter']['instruments']['pressure_gauge']['brand']) if
                self.electrical_specs['bagfilter']['instruments']['pressure_gauge']['brand'] else "")
            self.bagfilter_pt_brand.setCurrentText(
                str(self.electrical_specs['bagfilter']['instruments']['pressure_transmitter']['brand']) if
                self.electrical_specs['bagfilter']['instruments']['pressure_transmitter']['brand'] else "")
            self.bagfilter_ps_brand.setCurrentText(
                str(self.electrical_specs['bagfilter']['instruments']['pressure_switch']['brand']) if
                self.electrical_specs['bagfilter']['instruments']['pressure_switch']['brand'] else "")
            self.bagfilter_type.setCurrentText(
                str(self.electrical_specs['bagfilter']['type']) if self.electrical_specs['bagfilter']['type'] else "")
            self.bagfilter_outlet_tt_brand.setCurrentText(
                str(self.electrical_specs['bagfilter']['instruments']['outlet_temperature_transmitter']['brand']) if
                self.electrical_specs['bagfilter']['instruments']['outlet_temperature_transmitter']['brand'] else "")
            self.plc_series.setCurrentText(
                str(self.electrical_specs['bagfilter']['plc_series']) if self.electrical_specs['bagfilter'][
                    'plc_series'] else "")
            self.bagfilter_dpt_brand.setCurrentText(
                str(self.electrical_specs['bagfilter']['instruments']['delta_pressure_transmitter']['brand']) if
                self.electrical_specs['bagfilter']['instruments']['delta_pressure_transmitter']['brand'] else "")
            self.bagfilter_dps_brand.setCurrentText(
                str(self.electrical_specs['bagfilter']['instruments']['delta_pressure_switch']['brand']) if
                self.electrical_specs['bagfilter']['instruments']['delta_pressure_switch']['brand'] else "")
            self.touch_panel_model.setCurrentText(
                str(self.electrical_specs['bagfilter']['touch_panel']) if self.electrical_specs['bagfilter'][
                    'touch_panel'] else "None")
            self.plc_protocol.setCurrentText(
                str(self.electrical_specs['bagfilter']['plc_protocol']) if self.electrical_specs['bagfilter'][
                    'plc_protocol'] else "")
            self.bagfilter_inlet_tt_brand.setCurrentText(
                str(self.electrical_specs['bagfilter']['instruments']['inlet_temperature_transmitter']['brand']) if
                self.electrical_specs['bagfilter']['instruments']['inlet_temperature_transmitter']['brand'] else "")

            # QLineEdit elements
            self.fan_rpm.setText(str(self.electrical_specs['fan']['motors']['fan']['rpm']) if
                                 self.electrical_specs['fan']['motors']['fan']['rpm'] else "")
            self.bagfilter_order.setText(str(self.electrical_specs['bagfilter']['order']))

            # QCheckBox elements

            self.fan_checkbox.stateChanged.connect(self._handle_fan_checkbox_changed)
            self.fan_checkbox.setChecked(self.electrical_specs['fan']['status'])

            self.transport_checkbox.stateChanged.connect(self._handle_transport_checkbox_changed)
            self.transport_checkbox.setChecked(self.electrical_specs['transport']['status'])

            self.vibration_checkbox.stateChanged.connect(self._handle_vibration_checkbox_changed)
            self.vibration_checkbox.setChecked(self.electrical_specs['vibration']['status'])

            self.hopper_heater_checkbox.stateChanged.connect(self._handle_hopper_heater_checkbox_changed)
            self.hopper_heater_checkbox.setChecked(self.electrical_specs['hopper_heater']['status'])

            self.freshair_checkbox.stateChanged.connect(self._handle_freshair_checkbox_changed)
            self.freshair_checkbox.setChecked(self.electrical_specs['fresh_air']['status'])

            self.damper_checkbox.stateChanged.connect(self._handle_damper_checkbox_changed)
            self.damper_checkbox.setChecked(self.electrical_specs['damper']['status'])


            self.fan_space_heater.setChecked(self.electrical_specs['fan']['motors']['fan']['space_heater'])
            self.fan_de_nde.setChecked(self.electrical_specs['fan']['motors']['fan']['de_nde'])
            self.cable_supply.setChecked(self.electrical_specs['bagfilter']['cable_supply'])
            self.olm.setChecked(self.electrical_specs['bagfilter']['olm'])
            self.me.setChecked(self.electrical_specs['bagfilter']['me'])
            self.ee.setChecked(self.electrical_specs['bagfilter']['ee'])


        except KeyError as e:
            show_message(f"KeyError: Missing key in electrical_specs: {e}")
        except AttributeError as e:
            show_message(f"AttributeError: UI element not found: {e}")
        except Exception as e:
            show_message(f"Unexpected error: {e}")


def replace_placeholders(doc: Document, data: dict):
    def replace_in_paragraph(paragraph):
        original_text = ''.join(run.text for run in paragraph.runs)
        updated_text = original_text

        for key, value in data.items():
            updated_text = updated_text.replace(f"{{{{{key}}}}}", str(value))

        if updated_text != original_text:
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ''
            # Add the updated text in the first run or a new one
            if paragraph.runs:
                paragraph.runs[0].text = updated_text
            else:
                paragraph.add_run(updated_text)

    # Process all paragraphs in the main body
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    # Process all table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def create_qss_word():
    template_path = "assets/QSS-Template.docx"
    doc = Document(template_path)

    current_user = UserSession()
    current_project = ProjectSession()
    project_details = current_project.project_electrical_specs

    today_shamsi = jdatetime.datetime.today().strftime("%Y/%m/%d %H:%M")

    fan = project_details["fan"]["motors"]["fan"]
    anti_condensation_heater = "Yes (220vAc)" if fan["space_heater"] else "No"
    frequency_converter = "30Hz" if fan["start_type"] == "VFD" else "-"

    if project_details["fan"]["motors"]["fan"]["voltage_type"] == "LV":
        fan_voltage = project_details["project_info"]["l_voltage"]
    else:
        fan_voltage = project_details["project_info"]["m_voltage"]
    keywords = {
        "date": today_shamsi,
        "project_name": current_project.name,
        "project_code": current_project.code,
        "project_unique_code": current_project.unique_no,
        "humidity": project_details["project_info"]["humidity"],
        "min_temp": project_details["project_info"]["minimum_temprature"],
        "max_temp": project_details["project_info"]["maximum_temprature"],
        "altitude_elevation": project_details["project_info"]["altitude_elevation"],
        "power": fan["power"]/1000,
        "rpm": fan["rpm"],
        "voltage": fan_voltage,
        "voltage_variation": project_details["project_info"]["voltage_variation"],
        "frequency": project_details["project_info"]["voltage_frequency"],
        "frequency_variation": project_details["project_info"]["frequency_variation"],
        "brand": fan["brand"],
        "start_type": fan["start_type"],
        "cooling_method": fan["cooling_method"],
        "ip_rating": fan["ip_rating"],
        "efficiency_class": fan["efficiency_class"],
        "voltage_type": fan["voltage_type"],
        "painting_ral": fan["painting_ral"],
        "thermal_protection": fan["thermal_protection"],
        "anti_condensation_heater": anti_condensation_heater,
        "frequency_converter": frequency_converter,
        "check_user": f"{current_user.first_name[0].title()}.{current_user.last_name.title()}",
        "approve_user": f"{current_user.first_name[0].title()}.{current_user.last_name.title()}",
    }

    replace_placeholders(doc, keywords)

    output_path, _ = QFileDialog.getSaveFileName(
        caption="Save Word File",
        directory=f"QSS_For_Motor_{fan['power']/1000}KW_{fan['rpm']}RPM",
        filter="Word Documents (*.docx)"
    )

    if output_path:
        doc.save(output_path)
        os.startfile(output_path)
